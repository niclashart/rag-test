"""Document loaders for various file formats."""
from typing import List, Dict, Optional
from pathlib import Path
import docx
import pandas as pd
from bs4 import BeautifulSoup
import json
from logging_config.logger import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """Load documents from various file formats."""
    
    @staticmethod
    def load_txt(file_path: str) -> Dict[str, any]:
        """Load plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            "text": content,
            "metadata": {"file_type": "txt", "pages": 1}
        }
    
    @staticmethod
    def load_markdown(file_path: str) -> Dict[str, any]:
        """Load markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            "text": content,
            "metadata": {"file_type": "md", "pages": 1}
        }
    
    @staticmethod
    def load_docx(file_path: str) -> Dict[str, any]:
        """Load DOCX file."""
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        text = "\n".join(text_parts)
        
        return {
            "text": text,
            "metadata": {"file_type": "docx", "pages": len(doc.paragraphs)}
        }
    
    @staticmethod
    def load_csv(file_path: str) -> Dict[str, any]:
        """Load CSV file."""
        df = pd.read_csv(file_path)
        text = df.to_string(index=False)
        return {
            "text": text,
            "metadata": {"file_type": "csv", "rows": len(df), "columns": list(df.columns)}
        }
    
    @staticmethod
    def load_excel(file_path: str) -> Dict[str, any]:
        """Load Excel file."""
        df = pd.read_excel(file_path)
        text = df.to_string(index=False)
        return {
            "text": text,
            "metadata": {"file_type": "xlsx", "rows": len(df), "columns": list(df.columns)}
        }
    
    @staticmethod
    def load_html(file_path: str) -> Dict[str, any]:
        """Load HTML file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        soup = BeautifulSoup(html_content, 'html.parser')
        text = soup.get_text(separator='\n', strip=True)
        return {
            "text": text,
            "metadata": {"file_type": "html"}
        }
    
    @staticmethod
    def load_json(file_path: str) -> Dict[str, any]:
        """Load JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        text = json.dumps(data, indent=2, ensure_ascii=False)
        return {
            "text": text,
            "metadata": {"file_type": "json"}
        }
    
    @staticmethod
    def load_document(file_path: str) -> Dict[str, any]:
        """Load document based on file extension."""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        loader_map = {
            '.txt': DocumentLoader.load_txt,
            '.md': DocumentLoader.load_markdown,
            '.docx': DocumentLoader.load_docx,
            '.doc': DocumentLoader.load_docx,
            '.csv': DocumentLoader.load_csv,
            '.xlsx': DocumentLoader.load_excel,
            '.xls': DocumentLoader.load_excel,
            '.html': DocumentLoader.load_html,
            '.htm': DocumentLoader.load_html,
            '.json': DocumentLoader.load_json,
        }
        
        if extension == '.pdf':
            # PDF is handled separately in pdf_processor
            raise ValueError("Use PDFProcessor for PDF files")
        
        loader = loader_map.get(extension)
        if not loader:
            raise ValueError(f"Unsupported file type: {extension}")
        
        logger.info(f"Loading {extension} file: {file_path}")
        return loader(file_path)


